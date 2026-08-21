"""
Build AskMcNeese Final System Hard-Stoppage Development Word document.
Generates infographic PNGs, then assembles a professional .docx.
"""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm, Twips, Emu

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "AskMcNeese_Final_System_Hard_Stoppage_Development.docx"

# McNeese-inspired brand (professional blue + gold, not purple AI defaults)
NAVY = "#0B2D5C"
NAVY_RGB = (11, 45, 92)
GOLD = "#C4A35A"
GOLD_RGB = (196, 163, 90)
SKY = "#1E5A96"
LIGHT = "#F5F7FA"
SLATE = "#334155"
MUTED = "#64748B"
WHITE = "#FFFFFF"
GREEN = "#0F766E"
TEAL = "#0D9488"
AMBER = "#B45309"
ROSE = "#9F1239"


def _font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _rounded_rect(draw, xy, radius, fill, outline=None, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_cover_banner():
    w, h = 1600, 520
    img = Image.new("RGB", (w, h), NAVY_RGB)
    draw = ImageDraw.Draw(img)
    # diagonal gold accent band
    for i in range(0, w + h, 18):
        draw.line([(i, 0), (i - h, h)], fill=(20, 60, 110), width=1)
    draw.polygon([(0, h - 90), (w, h - 180), (w, h), (0, h)], fill=GOLD_RGB)
    draw.rectangle([(0, 0), (w, 8)], fill=GOLD_RGB)

    title = _font(54, bold=True)
    sub = _font(26)
    tiny = _font(18)
    draw.text((72, 90), "AskMcNeese", font=title, fill=WHITE)
    draw.text((72, 170), "Final System Hard-Stoppage Development", font=sub, fill=GOLD_RGB)
    draw.text(
        (72, 230),
        "Baseline freeze · Capability snapshot · Architecture · Roadmap to August 12 launch",
        font=tiny,
        fill=(200, 215, 235),
    )
    draw.text((72, 300), "McNeese ACM Student Chapter  ·  Summer 2026", font=tiny, fill=(180, 200, 225))
    draw.text((72, 350), "Reference branch: GitHub  ·  dev  ·  latest published commit Jul 10, 2026", font=tiny, fill=(160, 185, 210))
    img.save(ASSETS / "cover_banner.png", "PNG")


def make_system_today_infographic():
    w, h = 1400, 780
    img = Image.new("RGB", (w, h), (245, 247, 250))
    draw = ImageDraw.Draw(img)
    title = _font(32, bold=True)
    body = _font(18)
    small = _font(15)
    draw.text((48, 32), "What the system does today (plain English)", font=title, fill=NAVY_RGB)

    cards = [
        (48, 100, "1. Student asks", "A McNeese student types a\ncampus question in everyday\nlanguage — scholarships,\nhours, professors, programs."),
        (370, 100, "2. System searches", "It looks only in approved\npublic McNeese sources\n(saved knowledge + live\ncampus pages when allowed)."),
        (692, 100, "3. AI writes answer", "Claude drafts a clear reply\nusing those sources — not\nprivate grades, Canvas, or\nlogin-only student records."),
        (1014, 100, "4. Shows receipts", "The student sees citations\nand a live “what we’re\ndoing” trail while the\nanswer is being built."),
    ]
    for x, y, head, text in cards:
        _rounded_rect(draw, (x, y, x + 300, y + 260), 18, WHITE, NAVY_RGB, 2)
        draw.rectangle((x, y, x + 300, y + 48), fill=NAVY_RGB)
        draw.text((x + 18, y + 12), head, font=body, fill=WHITE)
        draw.multiline_text((x + 18, y + 70), text, font=small, fill=SLATE, spacing=6)

    # bottom strip
    _rounded_rect(draw, (48, 400, 1352, 720), 18, WHITE, GOLD_RGB, 3)
    draw.text((72, 424), "The simple promise", font=_font(24, bold=True), fill=NAVY_RGB)
    promise = (
        "AskMcNeese is a public campus Q&A assistant. It helps students find official McNeese information "
        "faster — with sources they can click — without inventing answers or touching private systems.\n\n"
        "Think of it as a smart campus librarian that only uses books (web pages/PDFs) the university "
        "already publishes for everyone."
    )
    draw.multiline_text((72, 480), promise, font=_font(17), fill=SLATE, spacing=8)
    img.save(ASSETS / "system_today.png", "PNG")


def make_pipeline_diagram():
    """DSA / data-system architecture flow."""
    fig, ax = plt.subplots(figsize=(14, 9.2), dpi=140)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9.2)
    ax.axis("off")
    fig.patch.set_facecolor("#F5F7FA")
    ax.set_facecolor("#F5F7FA")

    ax.text(
        7,
        8.85,
        "AskMcNeese — Program Data & System Architecture (DSA)",
        ha="center",
        va="center",
        fontsize=16,
        fontweight="bold",
        color=NAVY,
    )
    ax.text(
        7,
        8.4,
        "How a question moves from the student to a cited answer",
        ha="center",
        va="center",
        fontsize=11,
        color=MUTED,
    )

    def box(x, y, w, h, title, lines, fc=NAVY, ec=GOLD):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.15",
            facecolor=fc,
            edgecolor=ec,
            linewidth=1.8,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h - 0.28, title, ha="center", va="top", fontsize=10, fontweight="bold", color="white")
        for i, line in enumerate(lines):
            ax.text(x + w / 2, y + h - 0.55 - i * 0.28, line, ha="center", va="top", fontsize=8.5, color="#E2E8F0")

    def arrow(x1, y1, x2, y2, color=GOLD):
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color=color, lw=2),
        )

    # Top stack
    box(5.0, 7.2, 4.0, 0.9, "1 · Student (Browser)", ["React chat UI · no login yet"], fc="#0B2D5C")
    arrow(7, 7.2, 7, 6.85)
    box(4.3, 5.75, 5.4, 1.05, "2 · FastAPI  ·  POST /ask (SSE)", ["Accept question · live activity trail · stream answer"], fc="#1E5A96")
    arrow(7, 5.75, 7, 5.4)

    # Understand → RCCS → Answer (horizontal pipeline)
    box(0.5, 3.85, 3.5, 1.45, "3a · Understand", ["Intent / persona", "Clarify if needed", "Capability answers"], fc="#0F4C81")
    box(5.25, 3.85, 3.5, 1.45, "3b · RCCS Hybrid Plan", ["Classify question type", "Pick search channels", "Allowlist safety"], fc="#0D5C63")
    box(10.0, 3.85, 3.5, 1.45, "5 · Answer Writer", ["Claude grounded reply", "Structured sections", "Validated citations"], fc="#7C4A03")
    arrow(4.0, 4.55, 5.25, 4.55)
    # API drops into Understand
    ax.annotate(
        "",
        xy=(2.25, 5.3),
        xytext=(7, 5.75),
        arrowprops=dict(arrowstyle="-|>", color=GOLD, lw=1.8, connectionstyle="arc3,rad=-0.15"),
    )

    # Channels under RCCS only
    ax.text(7, 3.45, "4 · Selective retrieval channels (parallel when needed)", ha="center", fontsize=10, color=MUTED, fontweight="bold")
    box(0.4, 1.55, 3.0, 1.45, "Knowledge Base", ["ChromaDB chunks", "Crawled approved pages"], fc="#1E3A5F")
    box(3.8, 1.55, 3.0, 1.45, "Official Live Web", ["mcneese.edu now", "Registry + search"], fc="#1E3A5F")
    box(7.2, 1.55, 3.0, 1.45, "Companions", ["e.g. Rate My Professors", "Flag + registry gated"], fc="#1E3A5F")
    box(10.6, 1.55, 3.0, 1.45, "Agentic (Perplexity)", ["Sonar Pro research", "Domain-filtered search"], fc="#134E4A")

    for cx in [1.9, 5.3, 8.7, 12.1]:
        arrow(7.0, 3.85, cx, 3.05)

    # Evidence merge back to Answer Writer
    ax.annotate(
        "",
        xy=(11.75, 3.85),
        xytext=(7.0, 1.55),
        arrowprops=dict(arrowstyle="-|>", color="#C4A35A", lw=1.8, connectionstyle="arc3,rad=0.25"),
    )
    ax.text(9.6, 2.55, "evidence →", fontsize=8, color=GOLD, fontweight="bold")

    # Offline path
    box(
        0.4,
        0.15,
        13.2,
        1.05,
        "Offline ingestion (writer only)",
        ["knowledge/ registry  →  crawler fetch / clean / chunk / embed  →  ChromaDB   ·   Backend never writes the knowledge store"],
        fc="#334155",
        ec=GOLD,
    )

    fig.tight_layout()
    fig.savefig(ASSETS / "dsa_architecture.png", dpi=140, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def make_agentic_diagram():
    fig, ax = plt.subplots(figsize=(12, 6.5), dpi=140)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    fig.patch.set_facecolor("#F5F7FA")

    ax.text(6, 6.15, "Agentic Mode — what it actually is", ha="center", fontsize=15, fontweight="bold", color=NAVY)
    ax.text(
        6,
        5.7,
        "Not a free-roaming robot. A controlled research channel for live campus questions.",
        ha="center",
        fontsize=10,
        color=MUTED,
    )

    steps = [
        (0.4, 3.2, "Web mode ON", "User turns on\nWeb search"),
        (3.1, 3.2, "Plan says go", "RCCS decides\nofficial live\nis needed"),
        (5.8, 3.2, "Perplexity Sonar", "Domain-filtered\ncampus research\n+ citations"),
        (8.5, 3.2, "Page open (opt.)", "Fetch key pages\nfor fuller text"),
    ]
    for x, y, t, s in steps:
        patch = FancyBboxPatch((x, y), 2.4, 2.0, boxstyle="round,pad=0.02,rounding_size=0.12", facecolor=NAVY, edgecolor=GOLD, lw=1.6)
        ax.add_patch(patch)
        ax.text(x + 1.2, y + 1.5, t, ha="center", fontsize=10, fontweight="bold", color="white")
        ax.text(x + 1.2, y + 0.85, s, ha="center", fontsize=8.5, color="#CBD5E1")
    for x in [2.8, 5.5, 8.2]:
        ax.annotate("", xy=(x + 0.25, 4.2), xytext=(x - 0.05, 4.2), arrowprops=dict(arrowstyle="-|>", color=GOLD, lw=2))

    patch = FancyBboxPatch((0.4, 0.4), 11.2, 2.2, boxstyle="round,pad=0.02,rounding_size=0.12", facecolor="white", edgecolor=TEAL, lw=2)
    ax.add_patch(patch)
    ax.text(6, 2.2, "Safety rules baked in", ha="center", fontsize=11, fontweight="bold", color=NAVY)
    ax.text(
        6,
        1.3,
        "Only allowlisted / campus-relevant domains · Definitions stay knowledge-base first ·\n"
        "Companions & social are gated · Claude still writes the final student-facing answer ·\n"
        "Can be switched off with feature flags without breaking the core Ask path",
        ha="center",
        fontsize=9,
        color=SLATE,
    )
    fig.tight_layout()
    fig.savefig(ASSETS / "agentic_mode.png", dpi=140, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def make_progress_timeline():
    w, h = 1400, 720
    img = Image.new("RGB", (w, h), (245, 247, 250))
    draw = ImageDraw.Draw(img)
    draw.text((48, 28), "Development timeline — hard stoppage view", font=_font(30, bold=True), fill=NAVY_RGB)

    # timeline bar
    draw.rectangle((80, 160, 1320, 172), fill=GOLD_RGB)

    milestones = [
        (120, "GitHub baseline", "Jul 10\ndev commit", "Registry truth,\nchat shell,\n/ask pipeline"),
        (400, "RCCS hybrid", "Jul 12+", "Classify · plan ·\nsafe multi-channel\nretrieval"),
        (680, "Agentic + UX", "Jul 12–15", "Perplexity channel,\nlive activity,\nvisual overhaul"),
        (960, "Hard stoppage", "Jul 15", "Freeze snapshot\n+ roadmap for\nlaunch prep"),
        (1180, "Public launch", "Aug 12", "First release\nwindow"),
    ]
    for x, title, date, detail in milestones:
        draw.ellipse((x + 40, 148, x + 70, 178), fill=NAVY_RGB, outline=GOLD_RGB, width=3)
        _rounded_rect(draw, (x, 220, x + 200, 420), 14, WHITE, NAVY_RGB, 2)
        draw.text((x + 16, 240), title, font=_font(15, bold=True), fill=NAVY_RGB)
        draw.text((x + 16, 275), date, font=_font(14), fill=GOLD_RGB)
        draw.multiline_text((x + 16, 320), detail, font=_font(13), fill=SLATE, spacing=4)

    _rounded_rect(draw, (48, 480, 1352, 680), 16, NAVY_RGB)
    draw.text((72, 510), "How to read this document", font=_font(22, bold=True), fill=WHITE)
    draw.multiline_text(
        (72, 560),
        "Section A = what is already published on GitHub (dev).\n"
        "Section B = what the team built after that commit (local / in-flight).\n"
        "Section C = architecture (DSA).\n"
        "Section D = what we must build next to keep students hooked and grow ACM presence.",
        font=_font(16),
        fill=(210, 220, 235),
        spacing=6,
    )
    img.save(ASSETS / "timeline.png", "PNG")


def make_roadmap_infographic():
    w, h = 1400, 900
    img = Image.new("RGB", (w, h), (245, 247, 250))
    draw = ImageDraw.Draw(img)
    draw.text((48, 28), "Post–hard-stoppage roadmap", font=_font(30, bold=True), fill=NAVY_RGB)
    draw.text((48, 72), "Features that create retention, ACM visibility, and campus stickiness", font=_font(16), fill=MUTED)

    items = [
        (NAVY_RGB, "01", "Class Schedule", "Retention hook", "Students keep coming back because the schedule is a daily need. It solves a real McNeese student/professor pain and raises the stakes of using Ask + other campus tools."),
        ((13, 78, 99), "02", "ACM Member Dashboard", "Member workspace", "Secure member login and ACM-only work surfaces — meetings, tasks, chapter operations — not just a public brochure."),
        ((15, 92, 110), "03", "Public ACM Panel", "Separate subdomain", "A public information site for ACM (events, about, contact) living on its own subdomain, cleanly separated from member tools."),
        ((124, 83, 9), "04", "Tracking & SEO", "Web presence", "Google Analytics + Search Console (and SEO basics) so we can see who finds us, what they use, and how to improve discoverability."),
        ((127, 29, 42), "05", "Canvas Integration", "After Aug 12", "Connect after first launch. Private course systems stay out of scope until the public Ask product is live and stable."),
    ]
    y = 120
    for color, num, title, tag, blurb in items:
        _rounded_rect(draw, (48, y, 1352, y + 130), 16, WHITE, color, 3)
        draw.rectangle((48, y, 120, y + 130), fill=color)
        draw.text((68, y + 48), num, font=_font(28, bold=True), fill=WHITE)
        draw.text((150, y + 22), title, font=_font(22, bold=True), fill=NAVY_RGB)
        draw.text((150, y + 56), tag, font=_font(14, bold=True), fill=color)
        draw.multiline_text((150, y + 82), blurb, font=_font(14), fill=SLATE)
        y += 148

    img.save(ASSETS / "roadmap.png", "PNG")


def make_capability_matrix():
    fig, ax = plt.subplots(figsize=(12, 6.2), dpi=140)
    ax.axis("off")
    fig.patch.set_facecolor("#F5F7FA")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.2)
    ax.text(6, 5.85, "Capability snapshot at hard stoppage", ha="center", fontsize=14, fontweight="bold", color=NAVY)

    rows = [
        ("Public campus Q&A chat", "Ready"),
        ("Cited answers from approved sources", "Ready"),
        ("Knowledge base + live web modes", "Ready"),
        ("RCCS hybrid retrieval (flagged)", "Ready"),
        ("Agentic Perplexity research channel", "Ready"),
        ("Live activity / progress trail", "Ready"),
        ("World-class chat visual shell", "Ready"),
        ("ACM about / org command chain pages", "Ready"),
        ("Student login / class schedule", "Next"),
        ("ACM dashboards + public ACM subdomain", "Next"),
        ("SEO / Analytics / Search Console", "Next"),
        ("Canvas integration", "After Aug 12"),
    ]
    ax.add_patch(FancyBboxPatch((0.5, 0.4), 11, 5.1, boxstyle="round,pad=0.02,rounding_size=0.1", facecolor="white", edgecolor="#CBD5E1", lw=1))
    for i, (name, status) in enumerate(rows):
        y = 5.1 - i * 0.38
        ax.text(0.9, y, name, fontsize=10, color=SLATE, va="center")
        color = GREEN if status == "Ready" else (AMBER if status == "Next" else ROSE)
        ax.add_patch(FancyBboxPatch((9.0, y - 0.14), 2.0, 0.28, boxstyle="round,pad=0.01,rounding_size=0.08", facecolor=color, edgecolor=color))
        ax.text(10.0, y, status, ha="center", va="center", fontsize=8.5, color="white", fontweight="bold")
    fig.tight_layout()
    fig.savefig(ASSETS / "capability_matrix.png", dpi=140, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


# ---------------- Word helpers ----------------

def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 22, True, NAVY_RGB, "Calibri")
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, 16, True, NAVY_RGB, "Calibri")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, 13, True, (30, 90, 150), "Calibri")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, False, (51, 65, 85), "Calibri")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 11, True, NAVY_RGB)
        r2 = p.add_run(text)
        set_run_font(r2, 11, False, (51, 65, 85))
    else:
        r = p.add_run(text)
        set_run_font(r, 11, False, (51, 65, 85))
    p.paragraph_format.space_after = Pt(3)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=NAVY_RGB, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold, color)


def add_callout_table(doc, title, body, fill="0B2D5C"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    set_run_font(r1, 12, True, (255, 255, 255))
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, 10, False, (226, 232, 240))
    doc.add_paragraph()


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AskMcNeese  ·  Final System Hard-Stoppage Development  ·  McNeese ACM  ·  Confidential working draft  ·  July 15, 2026")
    set_run_font(run, 8, False, (100, 116, 139))


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_cover_banner()
    make_system_today_infographic()
    make_pipeline_diagram()
    make_agentic_diagram()
    make_progress_timeline()
    make_roadmap_infographic()
    make_capability_matrix()

    doc = Document()
    set_narrow_margins(doc)
    add_footer(doc)

    # Cover
    doc.add_picture(str(ASSETS / "cover_banner.png"), width=Inches(6.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document type: System freeze & roadmap brief for leadership / ACM / builders")
    set_run_font(r, 10, False, (100, 116, 139))

    add_callout_table(
        doc,
        "Hard stoppage meaning (in everyday words)",
        "This is the point where we say: “Here is what AskMcNeese already is,” freeze that picture for stakeholders, "
        "and clearly list what we still must build before and after the August 12 first launch. "
        "It is not a claim that coding stops forever — it is a clean baseline for the next chapter.",
    )

    # Section A
    add_heading_styled(doc, "A. What AskMcNeese is today (GitHub · dev)", 1)
    add_body(
        doc,
        "The latest commit published on the GitHub repository’s dev branch (as of this hard-stoppage snapshot) "
        "is from July 10, 2026. That published baseline is a working public campus Q&A prototype — not a full "
        "student portal yet.",
    )
    doc.add_picture(str(ASSETS / "system_today.png"), width=Inches(6.8))

    add_heading_styled(doc, "A1. One-paragraph product summary", 2)
    add_body(
        doc,
        "AskMcNeese is a campus AI chat helper for McNeese State University, built with the McNeese ACM Student Chapter. "
        "A student asks a normal question — scholarships, admissions, a professor, campus services — and the system "
        "looks through approved public McNeese web pages and documents, then uses Claude to write a clear answer with "
        "clickable sources. There is no student login in this baseline. Chat history stays in the browser. Private "
        "systems like Canvas grades or personal records are intentionally out of reach.",
    )

    add_heading_styled(doc, "A2. What already works on that published baseline", 2)
    add_bullet(doc, " Chat UI where students ask questions and read answers.", bold_lead="Public assistant:")
    add_bullet(doc, " Offline crawler that fetches approved pages/PDFs into a local knowledge store (ChromaDB).", bold_lead="Knowledge pipeline:")
    add_bullet(doc, " FastAPI backend that retrieves relevant pieces and asks Claude to answer with citations.", bold_lead="Ask API:")
    add_bullet(doc, " Approved URL registry so answers stay grounded in public McNeese information.", bold_lead="Source control:")
    add_bullet(doc, " Optional live mcneese.edu fetching in addition to the saved knowledge base.", bold_lead="Live web mode:")
    add_bullet(doc, " Health/status checks and basic query logging for operators.", bold_lead="Ops basics:")

    add_heading_styled(doc, "A3. What that baseline deliberately does not include", 2)
    add_bullet(doc, "No university SSO / student accounts.")
    add_bullet(doc, "No Canvas / private course data.")
    add_bullet(doc, "No ACM member dashboards or class schedule product.")
    add_bullet(doc, "No production hardening claim (CORS, hosting, analytics still open items).")

    # Section B
    add_heading_styled(doc, "B. What we built after the latest GitHub commit", 1)
    add_body(
        doc,
        "After the July 10 published commit, the team continued substantial work locally on the same dev line. "
        "This hard-stoppage brief treats that work as “completed / in hand” for planning, even where it is still "
        "awaiting a clean push to GitHub.",
    )
    doc.add_picture(str(ASSETS / "timeline.png"), width=Inches(6.8))

    add_heading_styled(doc, "B1. Smarter retrieval: RCCS hybrid", 2)
    add_body(
        doc,
        "We added a Retrieval Control & Channel Selection (RCCS) layer. In plain terms: the system first figures out "
        "what kind of question you asked, then chooses the right mix of places to look — saved campus knowledge, "
        "live official McNeese pages, and carefully gated companion sites (for example professor rating sources). "
        "It is selective and allowlist-safe, not “search the whole internet.”",
    )

    add_heading_styled(doc, "B2. Agentic mode (Perplexity)", 2)
    doc.add_picture(str(ASSETS / "agentic_mode.png"), width=Inches(6.6))
    add_body(
        doc,
        "Agentic mode is a controlled research helper. When web search is on and the plan needs fresh official "
        "information, Perplexity Sonar Pro can research across campus-relevant domains, return citations, and "
        "optionally open key pages for fuller text. Claude still writes the final student answer. Flags can turn "
        "this channel off without breaking core Ask.",
    )

    add_heading_styled(doc, "B3. Trust, activity, and answer quality", 2)
    add_bullet(doc, " Live “what we’re doing” trail (searching, found sources, writing) over SSE.", bold_lead="Activity events:")
    add_bullet(doc, " Clearer sectioned answers (facts, steps, warnings) when the model can extract them.", bold_lead="Structured answers:")
    add_bullet(doc, " Stronger citation validation against allowlists; prompt-injection hardening.", bold_lead="Citation integrity:")
    add_bullet(doc, " Honest replies when someone asks what the assistant can or cannot search.", bold_lead="Capability answers:")
    add_bullet(doc, " Optional supervisor orchestration path (default off) for deeper plan→execute flows.", bold_lead="Supervisor (optional):")

    add_heading_styled(doc, "B4. Product surface & visual system", 2)
    add_bullet(doc, " Major chat UI overhaul — glass shell, premium motion, mobile nav, composer docking.")
    add_bullet(doc, " System status, settings, and feedback panels.")
    add_bullet(doc, " About / team command-chain and related ACM-facing content pages.")
    add_bullet(doc, " Design system docs, visual validation records, and frontend test coverage growth.")

    add_heading_styled(doc, "B5. Engineering hygiene", 2)
    add_bullet(doc, " Large unit/test suite growth around RCCS, activity events, structured answers, and UI wiring.")
    add_bullet(doc, " Design records and implementation reports for auditability (RCCS, core stabilization, motion).")
    add_bullet(doc, " Companion source registry governance separate from official campus sources.")

    doc.add_picture(str(ASSETS / "capability_matrix.png"), width=Inches(6.5))

    # Section C
    add_heading_styled(doc, "C. DSA — Data & System Architecture", 1)
    add_body(
        doc,
        "DSA here means the program’s data and system architecture: the moving parts, who writes what, and how a "
        "student question becomes a cited answer. The diagram below is the hard-stoppage picture of the stack.",
    )
    doc.add_picture(str(ASSETS / "dsa_architecture.png"), width=Inches(6.8))

    add_heading_styled(doc, "C1. Layers in plain language", 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Layer", "Job", "Everyday analogy"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=(255, 255, 255), center=True)
        shade_cell(table.rows[0].cells[i], "0B2D5C")
    rows = [
        ("Frontend", "Chat screen students use", "The front desk"),
        ("API (/ask)", "Takes the question, streams progress + answer", "The dispatcher"),
        ("RCCS + retrieval", "Decides where to look and gathers evidence", "The research desk"),
        ("LLM (Claude)", "Writes the student-facing answer from evidence", "The librarian writing a brief"),
        ("Knowledge + crawler", "Offline: approved pages → searchable chunks", "The library stacks"),
    ]
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            set_cell_text(table.rows[r_i].cells[c_i], val, bold=(c_i == 0), size=9)
            if r_i % 2 == 0:
                shade_cell(table.rows[r_i].cells[c_i], "F1F5F9")
    doc.add_paragraph()

    add_heading_styled(doc, "C2. End-to-end question flow", 2)
    steps = [
        "Student types a question in the chat.",
        "Frontend sends it to POST /ask (streaming).",
        "Backend may clarify, detect persona, or answer capability questions.",
        "If RCCS is on: classify → plan → run selective channels in parallel.",
        "Evidence is cleaned, ranked, trust-separated, and citation-checked.",
        "Claude writes the answer; frontend shows markdown, citations, and activity trail.",
        "Conversation can stay in the browser (localStorage) — no server account yet.",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f" {s}", bold_lead=f"{i}.")

    add_heading_styled(doc, "C3. Non-negotiable safety idea", 2)
    add_callout_table(
        doc,
        "Approved public sources only",
        "The crawler is the only writer to the knowledge store. The backend reads at ask-time. "
        "Live and agentic channels still pass through allowlists / domain filters. "
        "Private student systems stay out until a later, deliberate integration (Canvas after August 12).",
        fill="134E4A",
    )

    # Mermaid as text for engineers who prefer source diagrams
    add_heading_styled(doc, "C4. Mermaid source (for engineering wikis)", 2)
    add_body(
        doc,
        "The visual above is the stakeholder view. Engineers can paste the following Mermaid into GitHub or Notion:",
        size=10,
    )
    mermaid = """flowchart TD
  U[Student question] --> FE[React chat UI]
  FE --> API[FastAPI POST /ask SSE]
  API --> UND[Understand: intent / persona / clarify]
  UND --> RCCS[RCCS classify + plan]
  RCCS --> KB[Knowledge Base Chroma]
  RCCS --> LIVE[Official live McNeese]
  RCCS --> COMP[Companions gated]
  RCCS --> AG[Agentic Perplexity]
  KB --> EV[Evidence merge + rank]
  LIVE --> EV
  COMP --> EV
  AG --> EV
  EV --> LLM[Claude grounded answer]
  LLM --> FE2[Citations + activity trail]
  REG[knowledge registry] --> CR[Crawler offline]
  CR --> KB"""
    p = doc.add_paragraph()
    run = p.add_run(mermaid)
    set_run_font(run, 8, False, (51, 65, 85), "Consolas")

    # Section D
    add_heading_styled(doc, "D. What we still have to work on", 1)
    add_body(
        doc,
        "These are the intentional next products after hard stoppage. They are chosen to keep students inside the "
        "McNeese ecosystem, grow ACM’s real operating surface, and measure whether people can find and use us.",
    )
    doc.add_picture(str(ASSETS / "roadmap.png"), width=Inches(6.8))

    add_heading_styled(doc, "D1. Class schedule — retention hook", 2)
    add_body(
        doc,
        "A class schedule feature is not a nice-to-have. It is the daily habit that keeps students returning. "
        "When AskMcNeese helps with “What’s my week look like?” / section times / professor-linked schedule needs, "
        "students stay in our McNeese ecosystem instead of bouncing to random tools. That higher daily use raises "
        "the stakes of discovering Ask, ACM, and other campus features. It should genuinely solve a McNeese "
        "student or professor problem — not feel like a gimmick.",
    )

    add_heading_styled(doc, "D2. ACM dashboards — member work", 2)
    add_body(
        doc,
        "ACM needs an internal dashboard with member login and member-only work: chapter operations, explicit ACM "
        "tasks, and the tools members actually use. This is separate from the public Ask chat.",
    )

    add_heading_styled(doc, "D3. Public ACM panel — different subdomain", 2)
    add_body(
        doc,
        "There should also be a public ACM information panel (about the chapter, events, how to join) living on "
        "its own subdomain. Public marketing/info and authenticated member work must not be mashed into one muddy "
        "screen. Subdomains make the separation obvious for students and for SEO.",
    )

    add_heading_styled(doc, "D4. Tracking & monitoring — SEO + analytics", 2)
    add_body(
        doc,
        "After we ship something people can visit, we must measure it. Plan: search-engine basics (SEO), Google "
        "Analytics for usage behavior, and Google Search Console for search presence and indexing health. Without "
        "this, we cannot tell if AskMcNeese or ACM pages are discoverable or useful in the wild.",
    )

    add_heading_styled(doc, "D5. Canvas integration — after August 12", 2)
    add_body(
        doc,
        "Canvas touches private course space. We intentionally sequence it after the first launch window "
        "(August 12). Launch the public, source-grounded Ask experience first; integrate Canvas only when the "
        "product, trust model, and access controls are ready.",
    )

    # Roadmap table
    add_heading_styled(doc, "D6. Priority table", 2)
    rt = doc.add_table(rows=6, cols=4)
    rt.style = "Table Grid"
    for i, h in enumerate(["Priority", "Workstream", "Why it matters", "Timing"]):
        set_cell_text(rt.rows[0].cells[i], h, bold=True, color=(255, 255, 255), center=True, size=9)
        shade_cell(rt.rows[0].cells[i], "0B2D5C")
    data = [
        ("P0", "Class schedule", "Daily retention / ecosystem hook", "Post hard-stop → launch prep"),
        ("P0", "ACM dashboards", "Member login + real chapter work", "Parallel with schedule"),
        ("P1", "Public ACM panel", "Separate subdomain for ACM info", "With ACM surface work"),
        ("P1", "SEO + Analytics + GSC", "Know usage & web presence", "At/after public deploy"),
        ("P2", "Canvas integration", "Private course systems", "After Aug 12 launch"),
    ]
    for r_i, row in enumerate(data, start=1):
        for c_i, val in enumerate(row):
            set_cell_text(rt.rows[r_i].cells[c_i], val, bold=(c_i == 0), size=9)
            if r_i % 2 == 0:
                shade_cell(rt.rows[r_i].cells[c_i], "F8FAFC")
    doc.add_paragraph()

    # Section E
    add_heading_styled(doc, "E. Hard-stoppage decision & next actions", 1)
    add_callout_table(
        doc,
        "Freeze statement",
        "As of July 15, 2026, AskMcNeese is declared at hard-stoppage for planning: a cited public campus Q&A "
        "system with hybrid retrieval and agentic research capability in hand, plus a modern chat product surface. "
        "The next development chapter is retention (class schedule), ACM dual surfaces (member + public subdomain), "
        "measurement (SEO/analytics), then Canvas after the August 12 first launch.",
    )

    add_heading_styled(doc, "Immediate recommended actions", 2)
    add_bullet(doc, " Publish the post–Jul-10 local stack to GitHub when review-ready (RCCS, agentic, UI).")
    add_bullet(doc, " Lock August 12 launch scope: public Ask + measurement; schedule/ACM can stage behind flags.")
    add_bullet(doc, " Start class-schedule discovery with students/professors (what pain is highest-stakes?).")
    add_bullet(doc, " Draft ACM subdomain + auth boundary (public panel vs member dashboard).")
    add_bullet(doc, " Prepare Analytics + Search Console property checklist before public DNS goes live.")
    add_bullet(doc, " Keep Canvas explicitly parked until after first launch.")

    add_heading_styled(doc, "Document control", 2)
    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    meta_rows = [
        ("Title", "AskMcNeese Final System Hard-Stoppage Development"),
        ("Baseline reference", "GitHub McNeeseACMChapter/askmcneese · branch dev · commit 11e81ca (Jul 10, 2026)"),
        ("In-hand advances", "RCCS hybrid, Perplexity agentic, activity SSE, visual overhaul, ACM about surfaces, tests/docs"),
        ("Audience", "ACM leadership, project advisor, PM, builders"),
        ("Date", "July 15, 2026"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        set_cell_text(meta.rows[i].cells[0], k, bold=True, size=9)
        shade_cell(meta.rows[i].cells[0], "E2E8F0")
        set_cell_text(meta.rows[i].cells[1], v, size=9)

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— End of hard-stoppage brief —")
    set_run_font(er, 10, True, GOLD_RGB)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_doc()
